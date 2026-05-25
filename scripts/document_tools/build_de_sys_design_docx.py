from __future__ import annotations

import re
import argparse
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/huangyining/Desktop/workspace/DE_sys")
MD = ROOT / "docs/design/数据集管理设计文档_V0.3.0_20260513.md"
DOCX = ROOT / "docs/design/数据集管理设计文档_V0.3.0_20260513.docx"
VERSION_LABEL = "V0.3.1"
DIAGRAM_DIR = Path("/private/tmp/de_sys_docx_diagrams")
DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)


def _font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
    ]
    for p in candidates:
        try:
            return ImageFont.truetype(p, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _wrap(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    text = text.replace("<br/>", " ").replace("<br>", " ")
    out: list[str] = []
    for raw in re.split(r"\s+", text):
        if not raw:
            continue
        cur = ""
        for ch in raw:
            trial = cur + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not cur:
                cur = trial
            else:
                out.append(cur)
                cur = ch
        if cur:
            if out and draw.textbbox((0, 0), out[-1] + " " + cur, font=font)[2] <= max_width:
                out[-1] += " " + cur
            else:
                out.append(cur)
    return out or [""]


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill=(59, 130, 246), width=3):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(y2 - y1) >= abs(x2 - x1):
        sign = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 7, y2 - sign * 12), (x2 + 7, y2 - sign * 12)]
    else:
        sign = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - sign * 12, y2 - 7), (x2 - sign * 12, y2 + 7)]
    draw.polygon(pts, fill=fill)


def _render_flowchart(code: str, out: Path) -> None:
    lines = [l.strip() for l in code.splitlines() if l.strip() and not l.strip().startswith("flowchart")]
    label_by_id: dict[str, str] = {}
    order: list[str] = []
    edges: list[tuple[str, str, str]] = []
    node_pat = re.compile(r"([A-Za-z][A-Za-z0-9_]*)\s*\[\s*\"([^\"]+)\"\s*\]")

    def add_node(node_id: str, label: str | None = None):
        if node_id not in label_by_id:
            order.append(node_id)
            label_by_id[node_id] = label or node_id
        elif label:
            label_by_id[node_id] = label

    for line in lines:
        for nid, label in node_pat.findall(line):
            add_node(nid, label)
        compact = node_pat.sub(lambda m: m.group(1), line)
        m = re.search(r"\b([A-Za-z][A-Za-z0-9_]*)\b\s*[-=.]+>\s*(?:\|([^|]+)\|)?\s*\b([A-Za-z][A-Za-z0-9_]*)\b", compact)
        if m:
            a, edge_label, b = m.group(1), m.group(2) or "", m.group(3)
            add_node(a)
            add_node(b)
            edges.append((a, b, edge_label))
    if not order:
        order = ["diagram"]
        label_by_id["diagram"] = code

    font = _font(24)
    small = _font(18)
    title = _font(20, True)
    box_w = 760
    box_h = 74
    gap = 44
    margin_x = 80
    margin_y = 60
    width = 920
    height = margin_y * 2 + len(order) * box_h + max(0, len(order) - 1) * gap
    img = Image.new("RGB", (width, max(260, height)), "white")
    draw = ImageDraw.Draw(img)
    positions: dict[str, tuple[int, int, int, int]] = {}
    y = margin_y
    for nid in order:
        x = margin_x
        rect = (x, y, x + box_w, y + box_h)
        positions[nid] = rect
        draw.rounded_rectangle(rect, radius=16, fill=(239, 246, 255), outline=(37, 99, 235), width=2)
        wrapped = _wrap(draw, label_by_id[nid], font, box_w - 36)
        total_h = len(wrapped) * 28
        ty = y + (box_h - total_h) // 2
        for line in wrapped[:3]:
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((x + (box_w - (bbox[2] - bbox[0])) / 2, ty), line, fill=(15, 23, 42), font=font)
            ty += 28
        y += box_h + gap
    for a, b, label in edges:
        if a not in positions or b not in positions:
            continue
        ax1, ay1, ax2, ay2 = positions[a]
        bx1, by1, bx2, by2 = positions[b]
        start = ((ax1 + ax2) // 2, ay2)
        end = ((bx1 + bx2) // 2, by1)
        _arrow(draw, start, end)
        # Mermaid branch labels are often short conditions; in a single-column
        # Word-friendly flow layout they tend to collide with boxes, so the
        # node labels carry the readable workflow and edge labels are omitted.
    img.save(out)


def _render_sequence(code: str, out: Path) -> None:
    lines = [l.strip() for l in code.splitlines() if l.strip() and not l.strip().startswith("sequenceDiagram")]
    participants: dict[str, str] = {}
    messages: list[tuple[str, str, str, bool]] = []
    for line in lines:
        m = re.match(r"participant\s+(\w+)\s+as\s+(.+)", line)
        if m:
            participants[m.group(1)] = m.group(2)
            continue
        m = re.match(r"(\w+)\s*(-{1,2}>>|->>)\s*(\w+)\s*:\s*(.+)", line)
        if m:
            a, arrow, b, label = m.groups()
            participants.setdefault(a, a)
            participants.setdefault(b, b)
            messages.append((a, b, label, arrow.startswith("--")))
    ids = list(participants.keys()) or ["FE", "BE"]
    font = _font(20)
    small = _font(17)
    width = max(900, 220 * len(ids) + 120)
    top = 70
    step = 72
    height = top + 80 + max(1, len(messages)) * step + 70
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    xs = {pid: int(80 + i * ((width - 160) / max(1, len(ids) - 1))) for i, pid in enumerate(ids)}
    for pid in ids:
        x = xs[pid]
        label = participants[pid]
        bw, bh = 150, 46
        draw.rounded_rectangle((x - bw // 2, top, x + bw // 2, top + bh), radius=12, fill=(239, 246, 255), outline=(37, 99, 235), width=2)
        wrapped = _wrap(draw, label, font, bw - 18)
        draw.text((x - bw // 2 + 10, top + 12), wrapped[0][:12], fill=(15, 23, 42), font=font)
        draw.line((x, top + bh, x, height - 45), fill=(203, 213, 225), width=2)
    y = top + 86
    for a, b, label, dashed in messages:
        x1, x2 = xs[a], xs[b]
        if dashed:
            seg = 10
            cur = min(x1, x2)
            end = max(x1, x2)
            while cur < end:
                draw.line((cur, y, min(cur + seg, end), y), fill=(59, 130, 246), width=3)
                cur += seg * 2
            _arrow(draw, (x2 - (1 if x2 > x1 else -1) * 14, y), (x2, y))
        else:
            _arrow(draw, (x1, y), (x2, y))
        lines_wrapped = _wrap(draw, label, small, max(170, abs(x2 - x1) - 30))
        tx = min(x1, x2) + 12
        for j, txt in enumerate(lines_wrapped[:2]):
            draw.text((tx, y - 25 + j * 20), txt, fill=(51, 65, 85), font=small)
        y += step
    img.save(out)


def render_mermaid(code: str, index: int) -> Path:
    out = DIAGRAM_DIR / f"diagram_{index:03d}.png"
    if code.lstrip().startswith("sequenceDiagram"):
        _render_sequence(code, out)
    else:
        _render_flowchart(code, out)
    return out


def set_font(run, name: str = "Calibri", east_asia: str = "Microsoft YaHei", size: float | None = None):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_table_fixed(table, widths: list[float]):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    grid_dxa = [int(w * 1440) for w in widths]
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(grid_dxa)))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w_dxa in grid_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w_dxa))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])


def split_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip().replace("\\|", "|") for c in s.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def compute_widths(rows: list[list[str]]) -> list[float]:
    n = max(len(r) for r in rows)
    weights = []
    for i in range(n):
        sample = [r[i] if i < len(r) else "" for r in rows[:8]]
        length = max(4, min(36, max(len(s) for s in sample)))
        weights.append(length)
    total = sum(weights)
    raw = [6.5 * w / total for w in weights]
    widths = [min(max(w, 0.55), 2.25) for w in raw]
    scale = 6.5 / sum(widths)
    return [round(w * scale, 3) for w in widths]


def add_inline_markdown(paragraph, text: str, base_size: float = 10.5, bold_default: bool = False):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Consolas", "Microsoft YaHei", base_size - 0.5)
            run.font.color.rgb = RGBColor(45, 55, 72)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size)
            run.bold = True
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size)
            run.bold = bold_default


def add_table(doc: Document, rows: list[list[str]]):
    if len(rows) >= 2 and is_table_separator("|" + "|".join(rows[1]) + "|"):
        rows = [rows[0]] + rows[2:]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    widths = compute_widths(rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_fixed(table, widths)
    font_size = 8.0 if ncols >= 6 else 8.7 if ncols >= 4 else 9.2
    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, text, font_size, bold_default=(r_idx == 0))
    doc.add_paragraph()


def build_docx(md_path: Path = MD, docx_path: Path = DOCX, version_label: str = VERSION_LABEL):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 20, "0B2545", 0, 10),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"数据集管理设计文档 {version_label}")
    set_font(run, size=8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lang = ""
    diagram_index = 0
    code_buf: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip().lower()
                code_buf = []
            else:
                if code_lang == "mermaid":
                    diagram_index += 1
                    img_path = render_mermaid("\n".join(code_buf), diagram_index)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(6.2))
                    doc.add_paragraph()
                else:
                    for code_line in code_buf:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.left_indent = Inches(0.18)
                        run = p.add_run(code_line if code_line else " ")
                        set_font(run, "Consolas", "Microsoft YaHei", 8.5)
                        run.font.color.rgb = RGBColor(31, 41, 55)
                    doc.add_paragraph()
                in_code = False
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, tbl_lines)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_inline_markdown(p, stripped[2:], 20, True)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(p, stripped[3:], 16, True)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, stripped[4:], 13, True)
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_markdown(p, stripped[5:], 12, True)
            i += 1
            continue
        if stripped.startswith("##### "):
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            add_inline_markdown(p, stripped[6:], 11, True)
            i += 1
            continue

        m_num = re.match(r"^(\d+)\\.\\s+(.*)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_inline_markdown(p, m_num.group(2), 10.5)
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline_markdown(p, stripped[2:], 10.5)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markdown(p, stripped, 10.5)
        i += 1

    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build DE_sys design DOCX from Markdown.")
    parser.add_argument("--input", type=Path, default=MD, help="Markdown source path")
    parser.add_argument("--output", type=Path, default=DOCX, help="DOCX output path")
    parser.add_argument("--version-label", default=VERSION_LABEL, help="Footer version label")
    args = parser.parse_args()
    build_docx(args.input, args.output, args.version_label)
